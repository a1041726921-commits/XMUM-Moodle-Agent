import json
import re
import urllib.error
import urllib.request
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Mapping, Optional

from .index import MaterialIndex
from .parser import parse_material


OUTPUT_DOCX_FILENAME = "Course_Knowledge_Checklist.docx"

COURSE_KNOWLEDGE_CHECKLIST_PROMPT = """你是一个“课程知识整理 Agent”，你的任务是根据用户提供的所有课件内容（包括 PDF、PPT、文本或 OCR 提取内容等），自动进行理解、整合与重构，并最终生成一份结构清晰、适合长期复习与理解的《课程知识清单》，输出为可下载的 .docx 文件。你需要先对所有输入内容进行整体阅读与语义理解，合并多份课件中的重复内容，统一术语表达，消除冗余信息，然后按照知识体系进行结构化重组。整理后的内容必须以“章节（Chapter/Module）→小节（Section）→知识点（Key Concepts）”的层级组织。每个关键知识点必须采用中英双语标注（中文 + English），以便国际学生或留学生理解，例如：线性变换（Linear Transformation）、特征值（Eigenvalue）、梯度下降（Gradient Descent）等。每个知识点下需包含：定义（Definition / 定义）、通俗解释（Explanation / 解释）、必要公式或规则（Formulas or Rules / 公式或规则，如果有）、以及典型例子（Examples / 示例）。对于复杂概念，需要补充直观理解或类比解释，以帮助用户真正理解其含义，而不是仅仅记忆。输出内容应当以“系统化知识整理”为目标，重点在于完整性、清晰性与可读性，而不是考试技巧或重点标记，因此不需要标注高频考点或重要程度等级。最终需要使用 Python 的 python-docx 库生成一个格式规范的 Word 文档，要求包含清晰的标题层级（标题、二级标题、三级标题）、合理的段落结构、适当加粗关键概念，并保证整体排版适合长期复习阅读。文件最终命名为 Course_Knowledge_Checklist.docx，并返回生成的文件或下载路径，不输出多余解释内容。"""

MAX_MATERIAL_CHARS = 12000
MAX_TOTAL_CHARS = 90000

JsonRequester = Callable[[str, Mapping[str, str], Mapping[str, object]], Mapping[str, object]]


def select_note_provider(providers: Iterable[Mapping[str, object]]) -> Optional[Dict[str, object]]:
    for provider in providers:
        if provider.get("enabled") and str(provider.get("api_key", "")).strip():
            return dict(provider)
    return None


def test_provider_connection(
    provider: Mapping[str, object],
    request_json: JsonRequester = None,
) -> str:
    request_json = request_json or _post_json
    url, headers, payload = build_provider_request(
        provider,
        "You are a connection test endpoint. Reply with only: ok",
        "ping",
        max_tokens=16,
    )
    response = request_json(url, headers, payload)
    text = _extract_response_text(response)
    return text or "ok"


def generate_course_knowledge_checklist(
    root: Path,
    provider: Mapping[str, object],
    request_json: JsonRequester = None,
) -> Path:
    root = root.resolve()
    index = MaterialIndex(root / "data" / "index.json")
    if not index.records:
        raise ValueError("没有找到已下载的课程资料，请先下载课程文件。")

    grouped_materials = _parse_index_materials(index)
    if not grouped_materials:
        raise ValueError("没有可读取的课件文本，请确认已下载 PDF、PPT、DOCX 或文本资料。")

    materials_prompt = build_materials_prompt(grouped_materials)
    generated_text = call_note_generation_api(provider, materials_prompt, request_json=request_json or _post_json)

    output_path = root / "data" / OUTPUT_DOCX_FILENAME
    write_ai_note_docx(generated_text, output_path)
    index.save()
    return output_path


def build_materials_prompt(grouped_materials: Mapping[str, List[Mapping[str, str]]]) -> str:
    parts: List[str] = [
        "本应用会使用 python-docx 将你的回答写入 Course_Knowledge_Checklist.docx；请直接输出文档正文内容。"
        "以下是用户提供的全部课程资料文本。请严格依据这些内容生成课程知识清单，不要编造课件中没有依据的具体事实。",
    ]
    total = 0
    for course_title, materials in grouped_materials.items():
        parts.append(f"\n\nCourse: {course_title}")
        for material in materials:
            text = str(material.get("text", "")).strip()
            if not text:
                continue
            remaining = MAX_TOTAL_CHARS - total
            if remaining <= 0:
                parts.append("\n[内容过长，后续材料已截断。]")
                return "".join(parts)
            clipped = text[: min(MAX_MATERIAL_CHARS, remaining)]
            total += len(clipped)
            parts.append(
                "\n\n"
                f"Material: {material.get('title', 'Untitled')}\n"
                f"Source: {material.get('source_file', '')}\n"
                f"{clipped}"
            )
    return "".join(parts)


def call_note_generation_api(
    provider: Mapping[str, object],
    materials_prompt: str,
    request_json: JsonRequester = None,
) -> str:
    request_json = request_json or _post_json
    url, headers, payload = build_provider_request(
        provider,
        COURSE_KNOWLEDGE_CHECKLIST_PROMPT,
        materials_prompt,
        max_tokens=8000,
        temperature=0.2,
    )
    response = request_json(url, headers, payload)
    return _extract_response_text(response)


def build_provider_request(
    provider: Mapping[str, object],
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = 8000,
    temperature: float = 0.2,
):
    provider_id = str(provider.get("id", "")).lower()
    base_url = str(provider.get("base_url", "")).strip().rstrip("/")
    api_key = str(provider.get("api_key", "")).strip()
    model = str(provider.get("model", "")).strip()
    if not base_url or not api_key or not model:
        raise ValueError("API 配置不完整，请检查提供方、API Key 和内置模型配置。")

    if provider_id == "anthropic":
        url = f"{base_url}/messages"
        headers = {
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        payload = {
            "model": model,
            "max_tokens": max_tokens,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        }
    else:
        url = f"{base_url}/chat/completions"
        headers = {
            "authorization": f"Bearer {api_key}",
            "content-type": "application/json",
        }
        payload = {
            "model": model,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
        }
    return url, headers, payload


def write_ai_note_docx(generated_text: str, output_path: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("缺少 python-docx，请先安装项目依赖后再生成 Word 笔记。") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    lines = [line.rstrip() for line in generated_text.splitlines()]
    wrote_content = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        heading = _markdown_heading(line)
        if heading:
            level, text = heading
            if level == 0:
                document.add_heading(text, level=0)
            else:
                document.add_heading(text, level=min(level, 3))
            wrote_content = True
            continue
        if line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_runs_with_bold(paragraph, line[2:].strip())
        else:
            paragraph = document.add_paragraph()
            _add_runs_with_bold(paragraph, line)
        wrote_content = True

    if not wrote_content:
        document.add_heading("课程知识清单", level=0)
        document.add_paragraph("AI 未返回可写入的笔记内容。")

    document.save(str(output_path))


def _parse_index_materials(index: MaterialIndex) -> Dict[str, List[Dict[str, str]]]:
    grouped: Dict[str, List[Dict[str, str]]] = {}
    for record in index.records:
        path = Path(record.local_path)
        if not path.exists():
            continue
        if not record.parsed_at or not record.text_excerpt:
            parsed = parse_material(path)
            index.mark_parsed(record, parsed)
        text = record.text_excerpt.strip()
        if not text:
            continue
        grouped.setdefault(record.course_title, []).append(
            {
                "title": record.parsed_title or record.resource_title,
                "source_file": record.local_path,
                "text": text,
            }
        )
    return grouped


def _post_json(url: str, headers: Mapping[str, str], payload: Mapping[str, object]) -> Mapping[str, object]:
    request = urllib.request.Request(
        url,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers=dict(headers),
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=120) as response:
            body = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        raise RuntimeError(f"AI API 请求失败：HTTP {exc.code} {detail}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"AI API 请求失败：{exc.reason}") from exc
    return json.loads(body)


def _extract_response_text(response: Mapping[str, object]) -> str:
    choices = response.get("choices")
    if isinstance(choices, list) and choices:
        first = choices[0]
        if isinstance(first, Mapping):
            message = first.get("message")
            if isinstance(message, Mapping) and message.get("content"):
                return str(message["content"]).strip()
            if first.get("text"):
                return str(first["text"]).strip()

    content = response.get("content")
    if isinstance(content, list):
        parts = []
        for item in content:
            if isinstance(item, Mapping) and item.get("text"):
                parts.append(str(item["text"]))
        if parts:
            return "\n".join(parts).strip()

    output_text = response.get("output_text")
    if output_text:
        return str(output_text).strip()

    raise RuntimeError("AI API 没有返回可用的笔记内容。")


def _markdown_heading(line: str):
    match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if not match:
        return None
    hashes, text = match.groups()
    level = len(hashes) - 1
    return level, text.strip()


def _add_runs_with_bold(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        run.bold = part.startswith("**") and part.endswith("**")
